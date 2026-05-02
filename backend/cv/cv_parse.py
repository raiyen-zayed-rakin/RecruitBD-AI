import json
import os
import sys

import docx
import pdfplumber
from ollama import chat

MODEL = "gemma3"


def extract_text(filepath):
    """Extract text from PDF or DOCX files."""
    if filepath.endswith(".pdf"):
        with pdfplumber.open(filepath) as pdf:
            return "\n".join(page.extract_text() for page in pdf.pages)
    elif filepath.endswith(".docx"):
        doc = docx.Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        return None


def parse_cv(text) -> dict:
    """Parse CV text into structured JSON according to the defined schema."""
    prompt = f"""
You are a strict JSON generator. Your task is to extract information from a CV.

RULES (MUST FOLLOW):
- Output ONLY valid JSON. No explanation, no markdown, no extra text.
- Do NOT add any fields that are not in the schema.
- Do NOT rename fields.
- Do NOT infer or guess missing information.
- If a field is missing, use:
  - "" for strings
  - [] for arrays
- Keep extracted text concise and factual.
- Do NOT include labels like "Name:", "Email:", etc.
- Do NOT hallucinate companies, dates, or skills.
- Use double quotes for JSON keys and string values.

SCHEMA:
{{
  "name": "",
  "email": "",
  "phone": [],
  "location": "",
  "summary": "",
  "skills": [],
  "experience": [
    {{
      "title": "",
      "company": "",
      "start_date": "",
      "end_date": "",
      "description": "",
      "tech": []
    }}
  ],
  "education": [
    {{
      "degree": "",
      "institution": "",
      "year": ""
    }}
  ],
  "languages": [],
  "certifications": [],
  "references": []
}}

EXTRACTION INSTRUCTIONS:
- Extract exactly what is written in the CV.
- For experience:
  - One object per role.
- For education:
  - One object per degree.
- For skills:
  - Extract explicit skills only (no assumptions).
- If no experience exist, return empty list [].
- If experience is a project, keep company as empty string and title as "[Project Name]".

CV TEXT:
{text}
"""


def extract_text(filepath):
    """Extract text from PDF or DOCX files."""
    if filepath.endswith(".pdf"):
        with pdfplumber.open(filepath) as pdf:
            return "\n".join(page.extract_text() for page in pdf.pages)
    elif filepath.endswith(".docx"):
        doc = docx.Document(filepath)
        text = []

        # Extract text from paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text.append(p.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for r in table.rows:
                row = []
                for cell in r.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row.append(cell_text)
                if row:
                    text.append(" | ".join(row))

        return "\n".join(text)
    else:
        return None


def parse_cv_ollama(text) -> dict:
    """Parse CV text into structured JSON according to the defined schema using Ollama."""
    prompt = PROMPT_TEMPLATE.format(text=text)

    response = chat(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": "You output ONLY JSON that strictly follows the schema. No variation allowed.",
            },
            {"role": "user", "content": prompt},
        ],
        format="json",
        think=False,
        options={
            "temperature": 0,
            "top_p": 1,
            "top_k": 1,
        },
    )

    if not response.done:
        raise ValueError("Model did not return a response.")

    raw = response.message.content
    if raw is None:
        raise ValueError("Model response is empty.")

    return json.loads(raw.strip())


def main():
    if len(sys.argv) != 2:
        print("Usage: python cv_parse.py <path_to_cv>")
        return

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"File {filepath} does not exist.")
        return

    text = extract_text(filepath)
    data = json.dumps(parse_cv(text))
    filename = os.path.splitext(os.path.basename(filepath))[0]
    output_path = f"{filename}_parsed.json"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(data)

    print(f"Output file: {output_path}")


if __name__ == "__main__":
    main()
